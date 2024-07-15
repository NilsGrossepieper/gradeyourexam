# Standard library imports
import io
import os
import re
import zipfile
from io import BytesIO

# External library imports
import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
import torch
from docx import Document
from docx.shared import RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from flask import Flask, render_template, request, redirect, url_for, send_file, session
from flask_session import Session
import requests
from sklearn.metrics.pairwise import cosine_similarity
from sklearn.preprocessing import PolynomialFeatures
from sklearn.linear_model import LinearRegression
from transformers import AutoTokenizer, AutoModel, AutoModelForCausalLM, GPT2Tokenizer

# Set matplotlib to use the 'Agg' backend to avoid issues with rendering in environments without display servers
plt.switch_backend('Agg')

# Initialize the Flask application
app = Flask(__name__)

# Load the model and tokenizer
API_URL = os.getenv("API_URL")
API_KEY = os.getenv("API_KEY")
headers = {
    "Accept": "application/json",
    "Authorization": f"Bearer {API_KEY}",
    "Content-Type": "application/json"
}

# Function to query the model API
def query_generative(payload):
    response = requests.post(API_URL, headers=headers, json=payload)
    return response.json()

# Load a tokenizer to count tokens (assuming you're using a model compatible with GPT-2 tokenizer)
tokenizer = GPT2Tokenizer.from_pretrained("gpt2")

def truncate_text(text, max_tokens):
    tokens = tokenizer.tokenize(text)
    if len(tokens) > max_tokens:
        tokens = tokens[:max_tokens]
    return tokenizer.convert_tokens_to_string(tokens)

# Set the session type to be stored in the filesystem
app.config['SESSION_TYPE'] = 'filesystem'

# Initialize the session for the Flask app
Session(app)

#################### / ####################

# Define the route for the main index page
@app.route('/', methods=['GET'])
def main_index():
    # Initialize session variables to store exam and evaluation data
    session['exam_data'] = []           # Stores data related to exams
    session['evaluation_data'] = []     # Stores data related to evaluations
    session['temp_df'] = []             # Temporary storage for dataframes
    session['grading_df'] = []          # Storage for grading dataframes
    session['solution_data'] = []       # Stores data related to solutions
    
    # Flags to indicate the success of file upload and regression analysis
    session['upload_success'] = False   # Indicates if file upload was successful
    session['regression_success'] = False  # Indicates if regression analysis was successful
    
    # Render the main index HTML template
    return render_template('main_index.html')

#################### /create_exam ####################

# Define the route for creating an exam, handling both GET and POST requests
@app.route('/create_exam', methods=['GET', 'POST'])
def create_exam_index():
    # Check if 'exam_data' is not in session, initialize it as an empty list
    if 'exam_data' not in session:
        session['exam_data'] = []

    # Handle POST request when a form is submitted
    if request.method == 'POST':
        # Retrieve the question, answer, and points from the form
        question = request.form.get('question')
        answer = request.form.get('answer')
        points = request.form.get('points')

        # Retrieve Alternative Solutions if defined in the form
        alternative_answer = request.form.get('alternative_answer')

        # Assign default value if Alternative Solutions are not provided
        if not alternative_answer:
            alternative_answer = ""

        # Create a new dictionary with the retrieved data
        new_data = {
            'Question': str(question),
            'Answer': str(answer),
            'Alternative Solution': str(alternative_answer),
            'Points': int(points)
        }

        # Retrieve the current exam data from the session
        exam_data = session['exam_data']
        # Append the new data to the exam data list
        exam_data.append(new_data)
        # Update the session with the modified exam data
        session['exam_data'] = exam_data

    # Convert the exam data in the session to a DataFrame
    exam_data_df = pd.DataFrame(session['exam_data'])
    # Convert the DataFrame to HTML for displaying in the template
    exam_data_html = exam_data_df.to_html(classes='data', header='true', index=False)
    
    # Render the 'create_exam_index.html' template, passing the HTML representation of the exam data
    return render_template('create_exam_index.html', exam_data_html=exam_data_html)

# Define the route for downloading the exam, handling POST requests
@app.route('/download_exam', methods=['POST'])
def download_exam():
    # Check if 'exam_data' is in the session, if not, redirect to the main index page
    if 'exam_data' not in session:
        return redirect(url_for('create_exam_index'))

    # Retrieve the exam data from the session and convert it to a DataFrame
    exam_data = pd.DataFrame(session['exam_data'])
    # Get the filename from the form data
    filename = request.form.get('filename')

    # Create new Word documents for the teacher's version and the student's version
    document_teacher = Document()
    document_student = Document()

    # Add a title to the teacher's document with centered alignment
    title_teacher_paragraph = document_teacher.add_heading(f'{filename} Sample Solution', level=1)
    title_teacher_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_teacher_run = title_teacher_paragraph.runs[0]
    title_teacher_run.bold = True
    title_teacher_run.underline = True
    title_teacher_run.font.color.rgb = RGBColor(0, 0, 0)

    # Add a title to the student's document with centered alignment
    title_student_paragraph = document_student.add_heading(f'{filename}', level=1)
    title_student_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_student_run = title_student_paragraph.runs[0]
    title_student_run.bold = True
    title_student_run.underline = True
    title_student_run.font.color.rgb = RGBColor(0, 0, 0)

    # Function to add student information fields to a document
    def add_student_info(document):
        document.add_paragraph('\n')
        paragraph_first_name = document.add_paragraph()
        hidden_first_name_before = paragraph_first_name.add_run('front_F1R5T_N4M3')
        hidden_first_name_before.font.hidden = True
        visible_first_name = paragraph_first_name.add_run('First name')
        visible_first_name.bold = True
        hidden_first_name_after = paragraph_first_name.add_run('end_F1R5T_N4M3')
        hidden_first_name_after.font.hidden = True
        visible_first_name_colon = paragraph_first_name.add_run(':')
        visible_first_name_colon.bold = True

        paragraph_surname = document.add_paragraph()
        hidden_surname_before = paragraph_surname.add_run('front_5URN4M3')
        hidden_surname_before.font.hidden = True
        visible_surname = paragraph_surname.add_run('Surname')
        visible_surname.bold = True
        hidden_surname_after = paragraph_surname.add_run('end_5URN4M3')
        hidden_surname_after.font.hidden = True
        visible_surname_colon = paragraph_surname.add_run(':')
        visible_surname_colon.bold = True

        paragraph_student_id = document.add_paragraph()
        hidden_student_id_before = paragraph_student_id.add_run('front_5TUD3NT_1D')
        hidden_student_id_before.font.hidden = True
        visible_student_id = paragraph_student_id.add_run('Student ID')
        visible_student_id.bold = True
        hidden_student_id_after = paragraph_student_id.add_run('end_5TUD3NT_1D')
        hidden_student_id_after.font.hidden = True
        visible_student_id_colon = paragraph_student_id.add_run(':')
        visible_student_id_colon.bold = True

        document.add_paragraph('\n')

    # Function to add exam questions and answers to the teacher's and student's documents
    def add_exam_questions_and_answers(document_teacher, document_student, exam_data):
        for index, row in exam_data.iterrows():
            
            # Add question details to the teacher's document
            paragraph_question_teacher = document_teacher.add_paragraph()
            hidden_question_teacher_before = paragraph_question_teacher.add_run(f'front_QU35T10N {index + 1}')
            hidden_question_teacher_before.font.hidden = True
            run_question_teacher = paragraph_question_teacher.add_run(f' Question {index + 1}:')
            run_question_teacher.bold = True
            hidden_question_teacher_after = paragraph_question_teacher.add_run(f'end_QU35T10N {index + 1}')
            hidden_question_teacher_after.font.hidden = True
            document_teacher.add_paragraph(row['Question'])

            paragraph_answer_teacher = document_teacher.add_paragraph()
            hidden_answer_teacher_before = paragraph_answer_teacher.add_run(f'front_4N5W3R {index + 1}')
            hidden_answer_teacher_before.font.hidden = True
            run_answer_teacher = paragraph_answer_teacher.add_run(f' Answer {index + 1}:')
            run_answer_teacher.bold = True
            hidden_answer_teacher_after = paragraph_answer_teacher.add_run(f'end_4N5W3R {index + 1}')
            hidden_answer_teacher_after.font.hidden = True
            document_teacher.add_paragraph(row['Answer'])
            
            paragraph_alternative_teacher = document_teacher.add_paragraph()
            hidden_alternative_teacher_before = paragraph_alternative_teacher.add_run(f'front_41T {index + 1}')
            hidden_alternative_teacher_before.font.hidden = True
            run_alternative_teacher = paragraph_alternative_teacher.add_run(f'Alternative Solution {index + 1}:')
            run_alternative_teacher.bold = True
            hidden_alternative_teacher_after = paragraph_alternative_teacher.add_run(f'end_41T {index + 1}')
            hidden_alternative_teacher_after.font.hidden = True
            document_teacher.add_paragraph(row['Alternative Solution'])

            paragraph_points_teacher = document_teacher.add_paragraph()
            hidden_points_teacher_before = paragraph_points_teacher.add_run(f'front_P01NT5 {index + 1}')
            hidden_points_teacher_before.font.hidden = True
            run_points_teacher = paragraph_points_teacher.add_run(f' Points for Question {index + 1}: ')
            run_points_teacher.bold = True
            hidden_points_teacher_after = paragraph_points_teacher.add_run(f'end_P01NT5 {index + 1}')
            hidden_points_teacher_after.font.hidden = True
            document_teacher.add_paragraph(str(row['Points']))

            document_teacher.add_paragraph('\n')

            # Add question details to the student's document
            paragraph_question_student = document_student.add_paragraph()
            hidden_question_student_before = paragraph_question_student.add_run(f'front_QU35T10N {index + 1}')
            hidden_question_student_before.font.hidden = True
            run_question_student = paragraph_question_student.add_run(f' Question {index + 1}: ')
            run_question_student.bold = True
            hidden_question_student_after = paragraph_question_student.add_run(f'end_QU35T10N {index + 1}')
            hidden_question_student_after.font.hidden = True
            document_student.add_paragraph(row['Question'])

            paragraph_answer_student = document_student.add_paragraph()
            hidden_answer_student_before = paragraph_answer_student.add_run(f'front_4N5W3R {index + 1}')
            hidden_answer_student_before.font.hidden = True
            run_answer_student = paragraph_answer_student.add_run(f' Answer {index + 1}: ')
            run_answer_student.bold = True
            hidden_answer_student_after = paragraph_answer_student.add_run(f'end_4N5W3R {index + 1}')
            hidden_answer_student_after.font.hidden = True
            document_student.add_paragraph('')

            paragraph_points_student = document_student.add_paragraph()
            hidden_points_student_before = paragraph_points_student.add_run(f'front_P01NT5 {index + 1}')
            hidden_points_student_before.font.hidden = True
            run_points_student = paragraph_points_student.add_run(f' Points for Question {index + 1}: ')
            run_points_student.bold = True
            hidden_points_student_after = paragraph_points_student.add_run(f'end_P01NT5 {index + 1}')
            hidden_points_student_after.font.hidden = True
            document_student.add_paragraph(str(row['Points']))

            document_student.add_paragraph('\n')

    # Add student info and exam questions/answers to both teacher and student documents
    add_student_info(document_teacher)
    add_student_info(document_student)
    add_exam_questions_and_answers(document_teacher, document_student, exam_data)

    # Save the teacher's document to a BytesIO object
    output_teacher = BytesIO()
    document_teacher.save(output_teacher)
    output_teacher.seek(0)

    # Save the student's document to a BytesIO object
    output_student = BytesIO()
    document_student.save(output_student)
    output_student.seek(0)

    # Create a zip file containing both documents
    zip_buffer = BytesIO()
    with zipfile.ZipFile(zip_buffer, 'w') as zip_file:
        zip_file.writestr(f'{filename}_sample_solution.docx', output_teacher.getvalue())
        zip_file.writestr(f'{filename}.docx', output_student.getvalue())
    zip_buffer.seek(0)

    # Send the zip file as a downloadable attachment
    return send_file(zip_buffer, as_attachment=True, download_name=f'{filename}.zip', mimetype='application/zip')

#################### /evaluate_exam ####################

# Define the route for evaluating exams, handling both GET and POST requests
@app.route('/evaluate_exam', methods=['GET', 'POST'])
def evaluate_exam_index():
    # Check if 'evaluation_data' is not in session, initialize it as an empty DataFrame
    if 'evaluation_data' not in session:
        session['evaluation_data'] = pd.DataFrame().to_dict()

    # Retrieve the evaluation data from the session and convert it to a DataFrame
    evaluation_data = pd.DataFrame(session['evaluation_data'])

    # Convert the evaluation data to HTML for rendering in the template
    if request.method == 'POST':
        sample_exam_file = request.files.get('sample_exam')
        student_exams = request.files.getlist('student_exams')

        # Check if both the sample exam file and student exams are provided
        if sample_exam_file and student_exams:
            sample_exam = Document(sample_exam_file)
            session['upload_success'] = True

            # Filter student solutions and sample solutions
            student_data, num_q_max, num_files = filter_student_solutions(student_exams)
            exam_data = filter_sample_solutions(sample_exam, num_q_max, num_files)
            evaluation_data = combine(exam_data, student_data) 
            session['evaluation_data'] = evaluation_data.to_dict()
            
        # If the sample exam file or student exams are not provided, set upload_success to False
        else:
            session['upload_success'] = False

    # Convert the evaluation data to HTML for rendering in the template
    evaluation_data_html_ready = pd.DataFrame(session['evaluation_data']).to_html(classes='data', header="true", index=False)
    return render_template('evaluate_exam_index.html', evaluation_data_html_ready=evaluation_data_html_ready, upload_success=session.get('upload_success', False))

# Function to filter sample solutions from the sample exam document
def filter_sample_solutions(doc, num_q, num_files):
    
    # Initialize sample_solution_data as an empty DataFrame
    sample_solution_data = pd.DataFrame(columns=['Question', 'Sample Solution', 'Alternative Solution', 'Points'])

    # Extract all text from the document
    full_text = []
    for paragraph in doc.paragraphs:
        full_text.append(paragraph.text)
    document_text = '\n'.join(full_text)

    # Extract sample solutions from the sample exam document
    for i in range(num_q):
        
        # Initialize question, answer, and points as None
        question = None
        answer = None
        points = None
        alternative = None

        # Extract question text using regex
        pattern_q = re.compile(rf'end_QU35T10N {i+1}(.*?)front_4N5W3R {i+1}', re.DOTALL)
        matches_q = pattern_q.finditer(document_text)
        for match_q in matches_q:
            question = match_q.group(1).strip()

        # Extract answer text using regex
        pattern_ss = re.compile(fr'end_4N5W3R {i+1}(.+?)front_41T {i+1}', re.DOTALL)
        matches_ss = pattern_ss.finditer(document_text)
        for match_ss in matches_ss:
            answer = match_ss.group(1).strip()
            
        # Extract alternativ answer 1 text using regex
        pattern_a = re.compile(fr'end_41T {i+1}(.+?)front_P01NT5 {i+1}', re.DOTALL)
        matches_a = pattern_a.finditer(document_text)
        for match_a in matches_a:
            alternative = match_a.group(1).strip()

        # Extract points using regex
        pattern_p = re.compile(rf'end_P01NT5 {i+1}\s*(\d+)', re.DOTALL)
        matches_p = pattern_p.finditer(document_text)
        for match_p in matches_p:
            points = match_p.group(1).strip()

        # Add the extracted data to the sample_solution_data DataFrame
        new_data_samp_sol = pd.DataFrame([{'Question': question, 'Sample Solution': answer,
                                        'Alternative Solution': alternative, 'Points': int(points)}])
        sample_solution_data = pd.concat([sample_solution_data, new_data_samp_sol], ignore_index=True)

    # Expand the sample_solution_data DataFrame to be as long as student_solution_data
    sample_solution_data = pd.concat([sample_solution_data]*num_files, ignore_index=True)

    return sample_solution_data

# Function to filter student solutions from the student exam documents
def filter_student_solutions(student_files):
    # Extract student details, questions, solutions, and points from student exam documents
    student_solution_data = pd.DataFrame(columns=['First Name', 'Surname', 'Student ID'])

    # Determine the number of files
    num_files = len(student_files)
    num_files = int(num_files)
    
    # List to store the number of questions in each student file
    num_q_list = []

    # Initialize student_solution_data_temp as an empty DataFrame
    for student_file in student_files:
        if student_file.filename.endswith('.docx'):
            full_text = []
            doc = Document(student_file)

            # Extract all text from the document
            for paragraph in doc.paragraphs:
                full_text.append(paragraph.text)
            document_text = '\n'.join(full_text)

            # Identify the number of questions using a regex pattern
            pattern_num_q = re.compile(r'front_QU35T10N .+? Question .+?: end_QU35T10N .+?', re.DOTALL)
            matches_num_q = pattern_num_q.finditer(document_text)
            num_q = int(len(list(matches_num_q)))
            num_q_list.append(num_q)

            student_solution_data_temp = pd.DataFrame(columns=['First Name', 'Surname', 'Student ID', 'Student Solution'])

            # Initialize first_name, surname, and student_id as None
            first_name = None
            surname = None
            student_id = None

            # Extract first name using regex
            pattern_fn = re.compile(r'end_F1R5T_N4M3(.*?)front_5URN4M3', re.DOTALL)
            matches_fn = pattern_fn.finditer(document_text)
            for match_fn in matches_fn:
                first_name = match_fn.group(1).replace(':', '').strip()

            # Extract surname using regex
            pattern_sn = re.compile(r'end_5URN4M3(.*?)front_5TUD3NT_1D', re.DOTALL)
            matches_sn = pattern_sn.finditer(document_text)
            for match_sn in matches_sn:
                surname = match_sn.group(1).replace(':', '').strip()

            # Extract student ID using regex
            pattern_id = re.compile(r'end_5TUD3NT_1D(.*?)front_QU35T10N 1', re.DOTALL)
            matches_id = pattern_id.finditer(document_text)
            for match_id in matches_id:
                student_id = match_id.group(1).replace(':', '').strip()

            for i in range(num_q):
                # Initialize student answer as None
                answer = None

                # Extract answer text using regex
                pattern_a = re.compile(rf'end_4N5W3R {i+1}(.*?)front_P01NT5 {i+1}', re.DOTALL)
                matches_a = pattern_a.finditer(document_text)
                for match_a in matches_a:
                    answer = match_a.group(1).strip()

                # Add the extracted data to the student_solution_data_temp DataFrame
                new_data_student_sol = pd.DataFrame([{'First Name': first_name, 'Surname': surname, 'Student ID': student_id, 'Student Solution': answer}])
                student_solution_data_temp = pd.concat([student_solution_data_temp, new_data_student_sol], ignore_index=True)

            # Concatenate the temporary DataFrame with the main student_solution_data DataFrame
            student_solution_data = pd.concat([student_solution_data, student_solution_data_temp], ignore_index=True)
            
    # Get the maximum number of questions in the student files 
    num_q_max = int(max(num_q_list))

    return student_solution_data, num_q_max, num_files

# Function to combine the sample exam data and student exam data for evaluation
def combine(exam_dataset, student_dataset):
    # Concate the exam_dataset and student_dataset DataFrames
    evaluation_data = pd.concat([student_dataset, exam_dataset], axis=1)    
    # Select relevant columns and rename 'Points_student' to 'Points'
    evaluation_data = evaluation_data[['First Name', 'Surname', 'Student ID', 'Question', 'Points', 'Sample Solution',
                                    'Alternative Solution', 'Student Solution']]
    
    solution_data = []
    
    # Get unique questions
    unique_questions = evaluation_data['Question'].unique()
    
    for question in unique_questions:
        # Get unique solutions for the current question
        unique_solutions = evaluation_data[evaluation_data['Question'] == question]['Sample Solution'].unique()
        
        # Append to the solution_data list as a dictionary
        solution_data.append({
            'Question': question,
            'Solution': list(unique_solutions)
        })
    
    # Update the solution_data in the session
    session['solution_data'] = solution_data

    # Clean exam_dataset and student_dataset
    exam_dataset = []
    student_dataset = []
    return evaluation_data

# Define the route for applying the model, handling POST requests
@app.route('/apply_model_llm', methods=['POST'])
def apply_model_llm():
    # If 'evaluation_data' is not in the session, redirect to the exam evaluation page
    if 'evaluation_data' not in session:
        return redirect(url_for('evaluate_exam_index'))

    # Retrieve evaluation data from the session and convert it to a DataFrame
    evaluation_data = pd.DataFrame(session['evaluation_data'])
    
    # Ensure all entries in the specified columns are strings
    evaluation_data['Student Solution'] = evaluation_data['Student Solution'].astype(str)
    evaluation_data['Alternative Solution'] = evaluation_data['Alternative Solution'].astype(str)
    evaluation_data['Sample Solution'] = evaluation_data['Sample Solution'].astype(str)
    
    # Separate rows with missing student solutions into a different DataFrame
    missing_answers = evaluation_data[evaluation_data['Student Solution'] == '']
    evaluation_data = evaluation_data[evaluation_data['Student Solution'] != '']
    
    # Prepare missing answers DataFrame for later use
    missing_answers['Solution Type'] = 'Sample Solution'
    missing_answers = missing_answers[['First Name', 'Surname', 'Student ID', 'Question',
                                    'Student Solution', 'Points', 'Solution Type', 'Sample Solution']]
    missing_answers.rename(columns={'Sample Solution': 'Solution'}, inplace=True)

    # Transform the evaluation_data DataFrame into long format for comparison
    evaluation_data = evaluation_data.melt(id_vars=['First Name', 'Surname', 'Student ID', 'Question',
                                                    'Student Solution', 'Points'],
                                            value_vars=['Sample Solution', 'Alternative Solution'],
                                            var_name='Solution Type', value_name='Solution')
    
    # Add a hypothetical cosine similarity column to the missing answers DataFrame
    missing_answers['Cosine Similarity'] = 0.0

    # Extract lists of student solutions and corresponding solutions for embedding
    student_data = evaluation_data['Student Solution'].tolist()
    sample_data = evaluation_data['Solution'].tolist()

    # Define a function to get embeddings for a list of texts
    def get_embeddings(texts):
        inputs = tokenizer(texts, padding=True, truncation=True, return_tensors='pt')
        with torch.no_grad():
            outputs = model(**inputs)
        embeddings = outputs.last_hidden_state.mean(dim=1)
        return embeddings

    # Load the tokenizer and model for generating embeddings
    tokenizer = AutoTokenizer.from_pretrained('GIST-all-MiniLM-L6-v2')
    model = AutoModel.from_pretrained('GIST-all-MiniLM-L6-v2')

    # Get embeddings for student solutions and sample solutions
    student_embeddings = get_embeddings(student_data)
    sample_embeddings = get_embeddings(sample_data)

    # Convert embeddings to numpy arrays for similarity calculation
    student_embeddings_array = student_embeddings.numpy()
    sample_embeddings_array = sample_embeddings.numpy()

    # Compute cosine similarity between student embeddings and sample embeddings
    similarity_matrix = cosine_similarity(student_embeddings_array, sample_embeddings_array)
    diagonal_similarities = np.diag(similarity_matrix)  # Get diagonal similarities (one-to-one)

    # Add cosine similarity values to the evaluation data DataFrame
    evaluation_data['Cosine Similarity'] = diagonal_similarities
    evaluation_data = pd.concat([evaluation_data, missing_answers], ignore_index=True)
    
    # Create a new Identification column for each student and question
    evaluation_data['Identification'] = evaluation_data.groupby(['Student ID', 'Question']).ngroup()
    
    # Filter out answers with the highest cosine similarity for each Identification group
    idx = evaluation_data.groupby('Identification')['Cosine Similarity'].idxmax()
    evaluation_data = evaluation_data.loc[idx]

    # Save the updated evaluation data back to the session
    session['evaluation_data'] = evaluation_data.to_dict()
    
    # Delete the tokenizer and model to free up memory
    del tokenizer
    del model
    
    # Redirect to the /llm route
    return redirect(url_for('llm'))

#################### /llm ####################

# Define the route for the LLM index page
@app.route('/llm', methods=['GET'])
def llm():
    # Ensure 'evaluation_data' is in the session
    if 'evaluation_data' not in session:
        return redirect(url_for('evaluate_exam_index'))
    # Return the llm_index.html template
    return render_template('llm_index.html')

# Define the route for applying the llm model, handling POST requests
@app.route('/apply_llm', methods=['POST'])
def apply_llm():
    # Ensure 'evaluation_data' is in the session
    if 'evaluation_data' not in session:
        return redirect(url_for('evaluate_exam_index'))
    
    # Retrieve evaluation data from the session and convert it to a DataFrame
    evaluation_data = pd.DataFrame(session['evaluation_data'])
    
    # Separate missing student solutions into a different DataFrame
    missing_answers_llm = evaluation_data[evaluation_data['Student Solution'] == '']
    evaluation_data = evaluation_data[evaluation_data['Student Solution'] != '']

    # Add two new columns to the evaluation data
    evaluation_data['Feedback'] = ''
    evaluation_data['Points for Student Answer'] = None  # Initialize as None
    missing_answers_llm['Feedback'] = 'Question has not been answered'
    missing_answers_llm['Points for Student Answer'] = 0
    missing_answers_llm['Student Solution'] = ''

    # Define the grading schemes
    grading_schemes = [
        'full points like 0, 1 or 2',
        'full points or half points like 0, 0.5 or 1',
        'full points or quarter points like 0, 0.25 or 0.5'
    ]

    # Extract naming scheme from POST request
    naming_scheme = request.json.get('naming_scheme')

    # Function to extract the first integer or float number from a string
    def extract_number(text):
        # Regular expression pattern to find the first number (integer or float)
        pattern = r"[-+]?\d*\.\d+|\d+"
        match = re.search(pattern, text)
        if match:
            return float(match.group())
        else:
            return None  # Return None if no number found

    # Iterate over each row in the evaluation data
    for index, row in evaluation_data.iterrows():
        # Determine which grading scheme to use based on the naming_scheme
        grading_scheme_to_use = grading_schemes[2]
        if naming_scheme == 'full_points':
            grading_scheme_to_use = grading_schemes[0]
        elif naming_scheme == 'half_points':
            grading_scheme_to_use = grading_schemes[1]

        # Create input text for the model using the selected grading scheme
        input_text = f"""You are a teacher grading exam questions, please provide a score from 0 to {row['Points']} for the following student answer. Use the following grading scheme: {grading_scheme_to_use}. Question: {row['Question']}, Student Solution: {row['Student Solution']}, Sample Solution: {row['Solution']}, Cosine Similarity between sample solution and student solution: {row['Cosine Similarity']}, Score:"""
        
        # Truncate the input text to ensure it does not exceed the token limit
        max_tokens = 924  # 1024 - 100 (to leave space for the model's response)
        input_text = truncate_text(input_text, max_tokens)
        
        response = query_generative({
            "inputs": input_text,
            "parameters": {
            "top_k": 3,
            "top_p": 0.5,
            "temperature": 0.5
        }})

        # Extract the generated text from the API response
        generated_text = response[0]['generated_text']
        # Extract the relevant part of the answer
        answer_start = generated_text.find(input_text) + len(input_text)
        answer = generated_text[answer_start:].strip()

        # Add the generated feedback to the evaluation data
        evaluation_data.at[index, 'Feedback'] = answer

        # Extract the first number from the generated text
        points_for_answer = extract_number(answer)

        # Store the extracted number in the 'Points for Student Answer' column
        evaluation_data.at[index, 'Points for Student Answer'] = points_for_answer
        
    # Concatenate the missing answers back to the evaluation data
    evaluation_data = pd.concat([evaluation_data, missing_answers_llm], ignore_index=True)
    
    # Sort by 'Surname', 'First Name' and 'Identification'
    evaluation_data = evaluation_data.sort_values(by=['Surname', 'First Name', 'Identification'])


    # Rearrange the columns as desired
    evaluation_data_export = evaluation_data[['First Name', 'Surname', 'Student ID', 'Question', 'Points', 'Solution', 'Student Solution', 'Points for Student Answer', 'Feedback']]

    # Create a BytesIO object to hold the Excel data
    excel_buffer = BytesIO()

    # Use pandas ExcelWriter to write the DataFrame to the BytesIO object
    with pd.ExcelWriter(excel_buffer, engine='xlsxwriter') as writer:
        evaluation_data_export.to_excel(writer, index=False, sheet_name='Sheet1')

    excel_buffer.seek(0)

    # Return the Excel file as a downloadable attachment
    return send_file(
        excel_buffer,
        as_attachment=True,
        download_name=f'llm_grading_{naming_scheme}.xlsx',
        mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
    )

# Define the route for applying the regression model, handling POST requests
@app.route('/apply_model_regression', methods=['POST'])
def apply_model_regression():
    # If 'evaluation_data' is not in the session, redirect to the exam evaluation page
    if 'evaluation_data' not in session:
        return redirect(url_for('evaluate_exam_index'))

    # Retrieve evaluation data from the session and convert it to a DataFrame
    evaluation_data = pd.DataFrame(session['evaluation_data'])
    
    # Ensure all entries in the specified columns are strings
    evaluation_data['Student Solution'] = evaluation_data['Student Solution'].astype(str)
    evaluation_data['Alternative Solution'] = evaluation_data['Alternative Solution'].astype(str)
    evaluation_data['Sample Solution'] = evaluation_data['Sample Solution'].astype(str)
    
    # Separate rows with missing student solutions into a different DataFrame
    missing_answers = evaluation_data[evaluation_data['Student Solution'] == ""]
    evaluation_data = evaluation_data[evaluation_data['Student Solution'] != ""]
    
    # Prepare missing answers DataFrame for later use
    missing_answers['Solution Type'] = 'Sample Solution'
    missing_answers = missing_answers[['First Name', 'Surname', 'Student ID', 'Question',
                                    'Student Solution', 'Points', 'Solution Type', 'Sample Solution']]
    missing_answers.rename(columns={'Sample Solution': 'Solution'}, inplace=True)

    # Transform the evaluation_data DataFrame into long format for comparison
    evaluation_data = evaluation_data.melt(id_vars=['First Name', 'Surname', 'Student ID', 'Question',
                                                    'Student Solution', 'Points'],
                                            value_vars=['Sample Solution', 'Alternative Solution'],
                                            var_name='Solution Type', value_name='Solution')
    
    # Add a hypothetical cosine similarity column to the missing answers DataFrame
    missing_answers['Cosine Similarity'] = 0.0

    # Extract lists of student solutions and corresponding solutions for embedding
    student_data = evaluation_data['Student Solution'].tolist()
    sample_data = evaluation_data['Solution'].tolist()

    # Define a function to get embeddings for a list of texts
    def get_embeddings(texts):
        inputs = tokenizer(texts, padding=True, truncation=True, return_tensors='pt')
        with torch.no_grad():
            outputs = model(**inputs)
        embeddings = outputs.last_hidden_state.mean(dim=1)
        return embeddings

    # Load the tokenizer and model for generating embeddings
    tokenizer = AutoTokenizer.from_pretrained('GIST-all-MiniLM-L6-v2')
    model = AutoModel.from_pretrained('GIST-all-MiniLM-L6-v2')

    # Get embeddings for student solutions and sample solutions
    student_embeddings = get_embeddings(student_data)
    sample_embeddings = get_embeddings(sample_data)

    # Convert embeddings to numpy arrays for similarity calculation
    student_embeddings_array = student_embeddings.numpy()
    sample_embeddings_array = sample_embeddings.numpy()

    # Compute cosine similarity between student embeddings and sample embeddings
    similarity_matrix = cosine_similarity(student_embeddings_array, sample_embeddings_array)
    diagonal_similarities = np.diag(similarity_matrix)  # Get diagonal similarities (one-to-one)

    # Add cosine similarity values to the evaluation data DataFrame
    evaluation_data['Cosine Similarity'] = diagonal_similarities
    evaluation_data = pd.concat([evaluation_data, missing_answers], ignore_index=True)
    
    # Create a new Identification column for each student and question
    evaluation_data['Identification'] = evaluation_data.groupby(['Student ID', 'Question']).ngroup()
    
    # Filter out answers with the highest cosine similarity for each Identification group
    idx = evaluation_data.groupby('Identification')['Cosine Similarity'].idxmax()
    evaluation_data = evaluation_data.loc[idx]
    
    # Load solution_data and evaluation_data from session
    solution_data = session['solution_data']
    
    # Convert solution_data to a dictionary for easy lookup
    solution_dict = {entry['Question']: entry['Solution'][0] if entry['Solution'] else '' for entry in solution_data}
    
    # Fill in empty entries in the Solution column
    for index, row in evaluation_data.iterrows():
        if row['Solution'] == '':
            question = row['Question']
            if question in solution_dict:
                evaluation_data.at[index, 'Solution'] = solution_dict[question]

    # Save the updated evaluation data back to the session
    session['evaluation_data'] = evaluation_data.to_dict()
    
    # Delete the tokenizer and model to free up memory
    del tokenizer
    del model

    # Initialize temporary DataFrames in session for grading
    session['temp_df'] = []
    session['grading_df'] = []

    # Add an empty 'Grading' column to the evaluation data
    evaluation_data['Grading'] = ""

    # Copy evaluation data to temp_df
    temp_df = evaluation_data.copy()

    # filter empty student solutions out of temp_df
    temp_df = temp_df[temp_df['Student Solution'] != '']
    
    # Split the data for grading based on the number of rows
    if temp_df.shape[0] < 5:
        grading_df = temp_df.copy()
        temp_df = temp_df.drop(grading_df.index)
        grading_df = grading_df.sample(frac=1).reset_index(drop=True)
    else:
        grading_df = temp_df.sample(5)
        temp_df = temp_df.drop(grading_df.index)
        
    # Select necessary columns for grading_df and temp_df
    grading_df = grading_df[['Identification', 'Question', 'Points', 'Solution', 'Student Solution', 'Grading']]
    temp_df = temp_df[['Identification', 'Question', 'Points', 'Solution', 'Student Solution', 'Grading']]

    # Save the temporary DataFrames back to the session
    session['temp_df'] = temp_df.to_dict()
    session['grading_df'] = grading_df.to_dict()

    # Redirect to the regression page for further processing
    return redirect(url_for('regression'))

# Define the route for the regression index page
@app.route('/regression', methods=['GET'])
def regression():
    # If 'evaluation_data' is not in the session, redirect to the exam evaluation page
    if 'evaluation_data' not in session:
        return redirect(url_for('evaluate_exam_index'))

    # Convert session data back to DataFrames
    evaluation_data = pd.DataFrame(session['evaluation_data'])
    temp_df = pd.DataFrame(session['temp_df'])
    grading_df = pd.DataFrame(session['grading_df'])

    # Generate HTML for the grading table
    grading_df_html_ready = generate_grading_table_html(grading_df)
    evaluation_data_html_ready = evaluation_data.to_html(classes='data', header='true', index=False)

    # Render the regression index template with the grading table
    return render_template('regression_index.html', grading_df_html_ready=grading_df_html_ready, evaluation_data_html_ready=evaluation_data_html_ready)

@app.route('/add_training_exercise', methods=['POST'])
def add_training_exercise():
    # Retrieve temporary and grading DataFrames from the session
    temp_df = pd.DataFrame(session.get('temp_df', []))
    grading_df = pd.DataFrame(session.get('grading_df', []))
    
    # Check if the temporary DataFrame is not empty
    if not temp_df.empty:
        # Ensure the DataFrame has the necessary columns
        temp_df = temp_df[['Identification', 'Question', 'Points', 'Solution', 'Student Solution', 'Grading']]
        
        # Select a random row from temp_df and remove it from temp_df
        new_row = temp_df.sample(1)
        temp_df = temp_df.drop(new_row.index)
        
        # Add the selected row to grading_df
        grading_df = pd.concat([grading_df, new_row])

        # Update session data with the modified DataFrames
        session['temp_df'] = temp_df.to_dict(orient='records')
        session['grading_df'] = grading_df.to_dict(orient='records')

    # Redirect to the regression page
    return redirect(url_for('regression'))

@app.route('/create_regression_table', methods=['POST'])
def create_regression_table():
    # Retrieve grading DataFrame from the session
    grading_df = pd.DataFrame(session.get('grading_df', []))
    
    # Update the 'Grading' column in the DataFrame with values from the form
    for idx, row in grading_df.iterrows():
        grading_df.at[idx, 'Grading'] = request.form.get(f'grading_{idx}', '')

    # Save the updated grading DataFrame back to the session
    session['grading_df'] = grading_df.to_dict(orient='records')

    # Call the create_regression function to process the grading DataFrame
    create_regression(grading_df)

    # Redirect to the regression page
    return redirect(url_for('regression'))

def generate_grading_table_html(grading_df):
    # Generate HTML for the grading table
    html = '<table class="data">'
    html += '<thead><tr>'
    
    # Add table headers
    for column in grading_df.columns:
        html += f'<th>{column}</th>'
    html += '</tr></thead><tbody>'
    
    # Add table rows
    for idx, row in grading_df.iterrows():
        html += '<tr>'
        for col in grading_df.columns:
            if col == 'Grading':
                max_value = row['Points']
                html += f'<td><input type="text" name="grading_{idx}" value="{row[col]}" max="{max_value}"></td>'
            else:
                html += f'<td>{row[col]}</td>'
        html += '</tr>'
    
    html += '</tbody></table>'
    
    return html

# Define the route for creating the regression model, handling POST requests
def create_regression(grading_df):
    # Retrieve the evaluation data from the session and merge with grading data
    evaluation_data = pd.DataFrame(session.get('evaluation_data', []))
    merged_df = grading_df.merge(evaluation_data, on='Identification', how='left')
    
    # Filter out rows where grading is empty or points are zero
    merged_df = merged_df[merged_df['Grading'] != ""]
    merged_df = merged_df[merged_df['Points_x'] != 0]
    
    # Calculate the relative points as a fraction of the total points
    merged_df['Relative Points'] = merged_df['Grading'].astype(float) / merged_df['Points_x'].astype(float)
    merged_df = merged_df[['Relative Points', 'Cosine Similarity']]
    
    # Add the points at extremes for better model fitting
    extreme_points = pd.DataFrame({
        'Relative Points': [0, 1],
        'Cosine Similarity': [0, 1]
    })
    merged_df = pd.concat([merged_df, extreme_points], ignore_index=True)

    # Prepare the data for regression
    X = merged_df[['Cosine Similarity']]
    y = merged_df['Relative Points']

    # Apply polynomial transformation (quadratic)
    poly = PolynomialFeatures(degree=2, include_bias=False)
    X_poly = poly.fit_transform(X)

    # Fit the linear regression model
    model = LinearRegression()
    model.fit(X_poly, y)
    
    # Extract the Cosine Similarity column from evaluation data for predictions
    X_eval = evaluation_data[['Cosine Similarity']]
    
    # Transform the Cosine Similarity values to polynomial features
    X_eval_poly = poly.transform(X_eval)
    
    # Predict the relative points using the trained model
    evaluation_data['Relative Points'] = model.predict(X_eval_poly)
    
    # Then, set Relative Points to 0 where Student Solution is empty
    evaluation_data.loc[evaluation_data['Student Solution'] == '', 'Relative Points'] = 0
    
    # Update the session with the modified evaluation data
    session['evaluation_data'] = evaluation_data.to_dict()
    
    # Indicate the regression process was successful
    session['regression_success'] = True

    # Generate predictions across the range 0 to 1 for plotting the fit
    x_range = np.linspace(0, 1, 100).reshape(-1, 1)
    x_range_poly = poly.transform(x_range)
    y_range_pred = model.predict(x_range_poly)

    # Plotting the data points and the quadratic fit
    plt.figure(figsize=(12.4, 6.2))  # Size in inches
    plt.scatter(merged_df['Cosine Similarity'], merged_df['Relative Points'], color='#F98A49', label='Data Points')
    plt.plot(x_range, y_range_pred, color='blue', label='Quadratic Fit')
    plt.xlabel('Cosine Similarity')
    plt.ylabel('Percentage of the Total Score')
    plt.title('Quadratic Model Fit')
    plt.legend()
    plt.grid(True)
    plt.xlim(0, 1)
    plt.ylim(0, 1)
    plt.yticks([0.0, 0.2, 0.4, 0.6, 0.8, 1.0], ['0%', '20%', '40%', '60%', '80%', '100%'])

    # Save the plot to a BytesIO object
    img = io.BytesIO()
    plt.savefig(img, format='png')
    img.seek(0)
    session['regression_plot'] = img.getvalue()

    # Close the plot to free up memory
    plt.close()

# Route to display the regression results
@app.route('/plot.png')
def plot_png():
    # Retrieve the regression plot image from the session
    img = session.get('regression_plot', None)
    
    # If the image is not found in the session, redirect to the regression page
    if img is None:
        return redirect(url_for('regression'))

    # Send the image file to the client
    return send_file(
        io.BytesIO(img),
        mimetype='image/png',
        as_attachment=False,  # Display the image in the browser
        download_name='plot.png'  # Set the default name for the image file
    )

# Endpoint to download the regression results as an Excel file
@app.route('/apply_download_regression', methods=['POST'])
def apply_download_regression():
    # Check if the regression process was successful
    if not session.get('regression_success', False):
        return redirect(url_for('regression'))

    # Retrieve the evaluation data from the session
    evaluation_data = pd.DataFrame(session['evaluation_data'])
    evaluation_data['Points for Student Answer'] = evaluation_data['Points'].astype(float) * evaluation_data['Relative Points'].astype(float)
    
    # Rearrange the columns as desired
    evaluation_data = evaluation_data[['Surname', 'First Name', 'Student ID', 'Question', 'Points', 'Solution', 'Student Solution', 'Points for Student Answer', 'Identification']]

    # Function to round grades based on the selected grading method
    def round_grades(grades, grading_method):
        rounded_grades = []
        for grade in grades:
            if grading_method == "Full Points":
                rounded_grade = np.ceil(grade)
            elif grading_method == "Half Points":
                rounded_grade = np.ceil(grade * 2) / 2
            elif grading_method == "Quarter Points":
                rounded_grade = np.ceil(grade * 4) / 4
            else:
                rounded_grade = grade
            rounded_grades.append(rounded_grade)
        return rounded_grades
    
    # Retrieve form data for grading method, naming scheme, and LLM explanation option
    grading_method = request.form.get('grading_method', 'Full Points')
    naming_scheme = request.form.get('naming_scheme', 'full_points')
    llm_true = request.form.get('llm_true', 'false') == 'true'

    # Perform rounding of grades
    evaluation_data_export_reg = evaluation_data.copy()
    evaluation_data_export_reg["Points for Student Answer"] = round_grades(evaluation_data["Points for Student Answer"], grading_method)

    # Optionally add feedback using a language model
    if llm_true:
        # Separate missing student solutions into a different DataFrame
        missing_answers_reg = evaluation_data_export_reg[evaluation_data_export_reg['Student Solution'] == '']
        evaluation_data_export_reg = evaluation_data_export_reg[evaluation_data_export_reg['Student Solution'] != '']
    
        missing_answers_reg['Feedback'] = 'Question has not been answered'
        missing_answers_reg['Points for Student Answer'] = 0
        
        evaluation_data_export_reg['Feedback'] = ''
        
        # Generate feedback for each student answer
        for index, row in evaluation_data_export_reg.iterrows():
            input_text = f"""You are a teacher grading exam questions, please explain the number of points a student has received based on the question, sample solution, and student answer. Provide short feedback why this student has received {row['Points for Student Answer']} points out of {row['Points']} total points for the following student answer. Question: {row['Question']}, Student Solution: {row['Student Solution']}, Sample Solution: {row['Solution']}, Feedback:"""
            
            # Truncate the input text to ensure it does not exceed the token limit
            max_tokens = 924  # 1024 - 100 (to leave space for the model's response)
            input_text = truncate_text(input_text, max_tokens)

            response = query_generative({
            "inputs": input_text,
            "parameters": {
            "top_k": 3,
            "top_p": 0.5,
            "temperature": 0.5
            }})

            # Extract the generated text from the API response
            generated_text = response[0]['generated_text']
            # Extract the relevant part of the answer
            answer_start = generated_text.find(input_text) + len(input_text)
            answer = generated_text[answer_start:].strip()

            # Add the generated feedback to the evaluation data
            evaluation_data_export_reg.at[index, 'Feedback'] = answer
        
        # Concatenate the missing answers back to the evaluation data
        evaluation_data_export_reg = pd.concat([evaluation_data_export_reg, missing_answers_reg], ignore_index=True)
        
    # Sort by 'Surname', 'First Name' and 'Identification'
    evaluation_data_export_reg = evaluation_data_export_reg.sort_values(by=['Surname', 'First Name', 'Identification'])
        
    # Remove the Identification column
    evaluation_data_export_reg = evaluation_data_export_reg.drop(columns=['Identification'])

    # Create an Excel buffer to store the data
    excel_buffer = BytesIO()

    # Write the DataFrame to an Excel file
    with pd.ExcelWriter(excel_buffer, engine='xlsxwriter') as writer:
        evaluation_data_export_reg.to_excel(writer, index=False, sheet_name='Sheet1')

    excel_buffer.seek(0)

    # Send the Excel file as a download
    return send_file(
        excel_buffer,
        as_attachment=True,
        download_name=f'regression_grading_{naming_scheme}.xlsx',
        mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
    )

# Main block to run the Flask application
if __name__ == '__main__':
    app.debug = True  # Enable debug mode for detailed error messages and live reloading
    port = int(os.environ.get('PORT', 5000))  # Get the port from environment variable or default to 5000
    app.run(host='0.0.0.0', port=port)  # Run the app on all available IP addresses (0.0.0.0) on the specified port